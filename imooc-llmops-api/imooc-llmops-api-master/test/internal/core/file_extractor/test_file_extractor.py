#!/usr/bin/env python
# -*- coding: utf-8 -*-
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document as DocxDocument

from internal.core.file_extractor import FileExtractor
from internal.entity.upload_file_entity import ALLOWED_DOCUMENT_EXTENSION


def test_load_from_file_docx_without_unstructured_runtime_download():
    with tempfile.TemporaryDirectory() as temp_dir:
        file_path = Path(temp_dir) / "sample.docx"

        doc = DocxDocument()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"
        table.cell(1, 0).text = "Value1"
        table.cell(1, 1).text = "Value2"
        doc.save(file_path)

        documents = FileExtractor.load_from_file(str(file_path), return_text=False, is_unstructured=True)

        assert len(documents) == 2
        assert "First paragraph" in documents[0].page_content
        assert "Second paragraph" in documents[0].page_content
        assert "Header1 | Header2" in documents[1].page_content
        assert "Value1 | Value2" in documents[1].page_content


def test_load_from_file_image_with_ocr():
    with tempfile.TemporaryDirectory() as temp_dir:
        file_path = Path(temp_dir) / "sample.png"

        image = Image.new("RGB", (500, 140), "white")
        drawer = ImageDraw.Draw(image)
        drawer.text((20, 40), "Hello OCR 123", fill="black")
        image.save(file_path)

        documents = FileExtractor.load_from_file(str(file_path), return_text=False, is_unstructured=True)

        assert len(documents) == 1
        assert "Hello OCR 123" in documents[0].page_content


def test_document_upload_extensions_include_common_ocr_images():
    for extension in ["png", "jpg", "jpeg", "webp", "gif", "bmp", "tif", "tiff"]:
        assert extension in ALLOWED_DOCUMENT_EXTENSION
