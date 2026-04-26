#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
@Time    : 2025/8/30 16:14
@Author  : ccckz@protonmail.com
@File    : file_extractor.py
"""
import os.path
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Union

import cv2
import requests
from docx import Document as DocxDocument
from injector import inject
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredCSVLoader,
    UnstructuredExcelLoader,
    UnstructuredFileLoader,
    UnstructuredHTMLLoader,
    UnstructuredMarkdownLoader,
    UnstructuredPowerPointLoader,
    UnstructuredXMLLoader,
)
from langchain_core.documents import Document as LCDocument
from rapidocr_onnxruntime import RapidOCR

from internal.entity.upload_file_entity import ALLOWED_DOCUMENT_IMAGE_EXTENSION
from internal.model import UploadFile
from internal.service import CosService


@inject
@dataclass
class FileExtractor:
    """文件提取器，用于将远程文件或 upload_file 记录加载为 LangChain 文档或文本。"""

    cos_service: CosService
    _ocr_engine = None

    def load(
        self,
        upload_file: UploadFile,
        return_text: bool = False,
        is_unstructured: bool = True,
    ) -> Union[list[LCDocument], str]:
        """加载传入的 upload_file 记录，返回 LangChain 文档列表或字符串。"""
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, os.path.basename(upload_file.key))
            self.cos_service.download_file(upload_file.key, file_path)
            return self.load_from_file(file_path, return_text, is_unstructured)

    @classmethod
    def load_from_url(cls, url: str, return_text: bool = False) -> Union[list[LCDocument], str]:
        """从 URL 下载并加载文档。"""
        response = requests.get(url)

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, os.path.basename(url))
            with open(file_path, "wb") as file:
                file.write(response.content)

            return cls.load_from_file(file_path, return_text)

    @classmethod
    def load_from_file(
        cls,
        file_path: str,
        return_text: bool = False,
        is_unstructured: bool = True,
    ) -> Union[list[LCDocument], str]:
        """从本地文件加载数据，返回 LangChain 文档列表或字符串。"""
        delimiter = "\n\n"
        file_extension = Path(file_path).suffix.lower()

        if file_extension in [".xlsx", ".xls"]:
            loader = UnstructuredExcelLoader(file_path)
        elif file_extension == ".pdf":
            loader = PyPDFLoader(file_path)
        elif file_extension == ".docx":
            documents = cls._load_docx(file_path)
            return delimiter.join([document.page_content for document in documents]) if return_text else documents
        elif file_extension.lstrip(".") in ALLOWED_DOCUMENT_IMAGE_EXTENSION:
            documents = cls._load_image_with_ocr(file_path)
            return delimiter.join([document.page_content for document in documents]) if return_text else documents
        elif file_extension in [".md", ".markdown"]:
            loader = UnstructuredMarkdownLoader(file_path)
        elif file_extension in [".htm", ".html"]:
            loader = UnstructuredHTMLLoader(file_path)
        elif file_extension == ".csv":
            loader = UnstructuredCSVLoader(file_path)
        elif file_extension in [".ppt", ".pptx"]:
            loader = UnstructuredPowerPointLoader(file_path)
        elif file_extension == ".xml":
            loader = UnstructuredXMLLoader(file_path)
        else:
            loader = UnstructuredFileLoader(file_path) if is_unstructured else TextLoader(file_path)

        documents = loader.load()
        return delimiter.join([document.page_content for document in documents]) if return_text else documents

    @classmethod
    def _load_docx(cls, file_path: str) -> list[LCDocument]:
        """使用 python-docx 本地解析，避免 unstructured 运行时下载 NLTK 资源。"""
        doc = DocxDocument(file_path)
        chunks: list[str] = []

        paragraphs = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text and paragraph.text.strip()
        ]
        if paragraphs:
            chunks.append("\n".join(paragraphs))

        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                chunks.append("\n".join(rows))

        if not chunks:
            return [LCDocument(page_content="", metadata={"source": file_path})]

        return [
            LCDocument(page_content=chunk, metadata={"source": file_path})
            for chunk in chunks
        ]

    @classmethod
    def _load_image_with_ocr(cls, file_path: str) -> list[LCDocument]:
        """使用 RapidOCR 本地提取图片文本，供知识库分段和索引使用。"""
        image = cv2.imread(file_path)
        if image is None:
            raise ValueError("图片文件读取失败，无法进行OCR识别")

        result, _ = cls._get_ocr_engine()(image)
        texts = [item[1].strip() for item in (result or []) if len(item) > 1 and item[1] and item[1].strip()]

        if not texts:
            return [LCDocument(page_content="", metadata={"source": file_path})]

        return [LCDocument(page_content="\n".join(texts), metadata={"source": file_path})]

    @classmethod
    def _get_ocr_engine(cls) -> RapidOCR:
        if cls._ocr_engine is None:
            cls._ocr_engine = RapidOCR()

        return cls._ocr_engine
